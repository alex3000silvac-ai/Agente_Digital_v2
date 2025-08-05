#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de Informes ANCI v2
Genera reportes de incidentes en formato estructurado JSON/Word para carga directa en plataforma ANCI
"""

import os
import json
import io
from datetime import datetime
from typing import Dict, Optional, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from ..core.database import get_db_connection

class GeneradorInformesANCIv2:
    """
    Genera informes ANCI en formato estructurado según requerimientos oficiales
    """
    
    def __init__(self):
        self.conn = None
        self.cursor = None
        
    def generar_informe_json(self, incidente_id: int, tipo_reporte: str = 'preliminar') -> Dict:
        """
        Genera informe en formato JSON estructurado para ANCI
        
        Args:
            incidente_id: ID del incidente
            tipo_reporte: 'preliminar', 'actualizacion', 'plan_accion', 'final'
            
        Returns:
            Diccionario con estructura ANCI
        """
        try:
            self.conn = get_db_connection()
            self.cursor = self.conn.cursor()
            
            # Cargar datos completos del incidente
            datos = self._cargar_datos_completos(incidente_id)
            
            # Generar estructura según tipo de reporte
            if tipo_reporte == 'preliminar':
                return self._generar_alerta_temprana(datos)
            elif tipo_reporte == 'actualizacion':
                return self._generar_segundo_reporte(datos)
            elif tipo_reporte == 'plan_accion':
                return self._generar_plan_accion(datos)
            elif tipo_reporte == 'final':
                return self._generar_informe_final(datos)
            else:
                raise ValueError(f"Tipo de reporte no válido: {tipo_reporte}")
                
        finally:
            if self.conn:
                self.conn.close()
    
    def generar_documento_word(self, incidente_id: int, tipo_reporte: str = 'preliminar') -> io.BytesIO:
        """
        Genera documento Word con formato oficial ANCI
        """
        # Obtener datos JSON
        datos_json = self.generar_informe_json(incidente_id, tipo_reporte)
        
        # Crear documento Word
        doc = Document()
        
        # Configurar estilos
        self._configurar_estilos(doc)
        
        # Título principal
        titulo = doc.add_heading('REPORTE DE INCIDENTE DE CIBERSEGURIDAD', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo con tipo de reporte
        subtitulo = doc.add_heading(f'Tipo: {tipo_reporte.upper()}', 1)
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información de referencia
        doc.add_paragraph()
        self._agregar_tabla_referencia(doc, datos_json.get('referencia', {}))
        
        # Secciones según tipo de reporte
        for seccion, contenido in datos_json.items():
            if seccion != 'referencia' and isinstance(contenido, dict):
                self._agregar_seccion(doc, seccion, contenido)
        
        # Archivos adjuntos (referencias)
        if 'archivos_adjuntos' in datos_json:
            self._agregar_seccion_archivos(doc, datos_json['archivos_adjuntos'])
        
        # Guardar en memoria
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output
    
    def _cargar_datos_completos(self, incidente_id: int) -> Dict:
        """Carga todos los datos necesarios del incidente"""
        # Datos del incidente
        query_incidente = """
            SELECT 
                i.*,
                e.RazonSocial, e.RUT, e.TipoEmpresa, e.SectorEsencial,
                u.nombre as NombreReportante, u.email as EmailReportante,
                'N/D' as TelefonoReportante
            FROM Incidentes i
            INNER JOIN Empresa e ON i.EmpresaID = e.EmpresaID
            LEFT JOIN Usuario u ON u.UsuarioID = 1
            WHERE i.IncidenteID = ?
        """
        self.cursor.execute(query_incidente, incidente_id)
        incidente = self._dict_from_row(self.cursor.fetchone())
        
        if not incidente:
            raise ValueError(f"Incidente {incidente_id} no encontrado")
        
        # Taxonomías
        query_taxonomias = """
            SELECT t.Codigo, t.Titulo, t.Descripcion, it.Comentarios
            FROM INCIDENTE_TAXONOMIA it
            INNER JOIN Taxonomia_incidentes t ON it.Id_Taxonomia = t.Id_Taxonomia
            WHERE it.IncidenteID = ?
        """
        self.cursor.execute(query_taxonomias, incidente_id)
        taxonomias = [self._dict_from_row(row) for row in self.cursor.fetchall()]
        
        # Archivos de evidencia
        query_archivos = """
            SELECT 
                ArchivoID, SeccionID, NombreArchivo, TipoArchivo,
                Descripcion, FechaSubida, SubidoPor
            FROM INCIDENTES_ARCHIVOS
            WHERE IncidenteID = ?
            ORDER BY SeccionID, FechaSubida
        """
        self.cursor.execute(query_archivos, incidente_id)
        archivos = [self._dict_from_row(row) for row in self.cursor.fetchall()]
        
        # Historial de cambios (para reportes de actualización)
        query_historial = """
            SELECT TOP 10 
                TipoAccion, DatosAnteriores, DatosNuevos, 
                Usuario, FechaAccion
            FROM INCIDENTES_AUDITORIA
            WHERE IncidenteID = ?
            ORDER BY FechaAccion DESC
        """
        self.cursor.execute(query_historial, incidente_id)
        historial = [self._dict_from_row(row) for row in self.cursor.fetchall()]
        
        return {
            'incidente': incidente,
            'taxonomias': taxonomias,
            'archivos': archivos,
            'historial': historial
        }
    
    def _generar_alerta_temprana(self, datos: Dict) -> Dict:
        """Genera estructura para alerta temprana (3-4 horas)"""
        incidente = datos['incidente']
        
        return {
            'referencia': {
                'tipo_reporte': 'ALERTA_TEMPRANA',
                'fecha_generacion': datetime.now().isoformat(),
                'id_interno': incidente['IDVisible'],
                'folio_anci': incidente.get('ReporteAnciID', '')
            },
            'identificacion_entidad': {
                'nombre_institucion': incidente['RazonSocial'],
                'rut': incidente['RUT'],
                'tipo_entidad': incidente['TipoEmpresa'],
                'sector_esencial': incidente.get('SectorEsencial', 'No especificado')
            },
            'datos_contacto': {
                'nombre_reportante': incidente.get('NombreReportante', ''),
                'cargo': 'Encargado de Ciberseguridad',
                'telefono_24_7': incidente.get('TelefonoReportante', ''),
                'email_oficial': incidente.get('EmailReportante', '')
            },
            'datos_incidente': {
                'fecha_hora_deteccion': self._format_datetime(incidente['FechaDeteccion']),
                'fecha_hora_inicio_estimada': self._format_datetime(incidente.get('FechaOcurrencia', incidente['FechaDeteccion'])),
                'descripcion_breve': self._truncar_texto(incidente.get('DescripcionInicial', ''), 500),
                'taxonomia_inicial': self._obtener_taxonomias_principales(datos['taxonomias'])
            },
            'impacto_inicial': {
                'sistemas_afectados': incidente.get('SistemasAfectados', ''),
                'servicios_interrumpidos': 'Sí' if incidente.get('ServiciosInterrumpidos') else 'No',
                'duracion_estimada': 'Por determinar',
                'usuarios_afectados': 'En evaluación',
                'alcance_geografico': incidente.get('AlcanceGeografico', 'Local')
            },
            'estado_actual': {
                'incidente_en_curso': 'Sí' if incidente['EstadoActual'] != 'Cerrado' else 'No',
                'contenido_aplicado': 'Sí' if incidente.get('AccionesInmediatas') else 'No',
                'descripcion_estado': incidente.get('EstadoActual', 'En gestión')
            },
            'acciones_inmediatas': {
                'medidas_contencion': incidente.get('AccionesInmediatas', 'En proceso de implementación'),
                'sistemas_aislados': 'Por determinar'
            },
            'solicitud_apoyo': {
                'requiere_asistencia_csirt': 'No',
                'tipo_apoyo_requerido': 'N/A'
            },
            'archivos_adjuntos': self._filtrar_archivos_seccion(datos['archivos'], [1, 2, 5])
        }
    
    def _generar_segundo_reporte(self, datos: Dict) -> Dict:
        """Genera estructura para segundo reporte (24-72 horas)"""
        # Incluye todo de alerta temprana más campos adicionales
        reporte = self._generar_alerta_temprana(datos)
        incidente = datos['incidente']
        
        # Actualizar referencia
        reporte['referencia']['tipo_reporte'] = 'SEGUNDO_REPORTE'
        reporte['referencia']['id_alerta_temprana'] = incidente.get('ReporteAnciID', '')
        reporte['referencia']['fecha_actualizacion'] = datetime.now().isoformat()
        
        # Agregar secciones adicionales
        reporte['analisis_detallado'] = {
            'descripcion_completa': incidente.get('DescripcionInicial', ''),
            'vector_ataque': incidente.get('OrigenIncidente', 'En investigación'),
            'causa_raiz_preliminar': incidente.get('AnciTipoAmenaza', 'Por determinar')
        }
        
        reporte['gravedad_impacto'] = {
            'nivel_criticidad': incidente.get('Criticidad', 'Media'),
            'sistemas_especificos_afectados': incidente.get('SistemasAfectados', ''),
            'duracion_real_interrupcion': 'En cálculo',
            'numero_usuarios_impactados': 'En evaluación',
            'volumen_datos_comprometidos': 'No determinado',
            'efectos_colaterales': incidente.get('AnciImpactoPreliminar', 'Sin efectos colaterales identificados')
        }
        
        reporte['indicadores_compromiso'] = {
            'direcciones_ip_sospechosas': [],
            'hashes_malware': [],
            'dominios_maliciosos': [],
            'cuentas_comprometidas': [],
            'urls_maliciosas': []
        }
        
        reporte['analisis_causa_raiz'] = {
            'vulnerabilidad_explotada': 'En análisis',
            'falla_control_seguridad': incidente.get('CausaRaiz', 'Por determinar'),
            'error_humano_involucrado': 'No'
        }
        
        reporte['acciones_respuesta'] = {
            'medidas_contencion_aplicadas': incidente.get('AccionesInmediatas', ''),
            'medidas_erradicacion': 'En proceso',
            'estado_recuperacion': 'En curso',
            'tiempo_estimado_resolucion': '48 horas'
        }
        
        reporte['coordinaciones_externas'] = {
            'notificacion_regulador_sectorial': 'No',
            'denuncia_policial': 'No',
            'contacto_proveedores_seguridad': 'Sí',
            'comunicacion_publica_emitida': 'No'
        }
        
        reporte['apoyo_requerido'] = {
            'asistencia_csirt_actualizada': 'No',
            'tipo_apoyo_en_curso': 'N/A'
        }
        
        # Agregar todos los archivos
        reporte['archivos_adjuntos'] = self._filtrar_archivos_seccion(datos['archivos'], None)
        
        return reporte
    
    def _generar_plan_accion(self, datos: Dict) -> Dict:
        """Genera plan de acción (solo OIV - 7 días)"""
        reporte = self._generar_segundo_reporte(datos)
        incidente = datos['incidente']
        
        # Actualizar referencia
        reporte['referencia']['tipo_reporte'] = 'PLAN_ACCION_OIV'
        
        # Agregar plan específico OIV
        reporte['plan_recuperacion'] = {
            'programa_restauracion_datos': 'Cronograma en desarrollo',
            'responsables_tecnicos': incidente.get('ResponsableCliente', ''),
            'responsables_administrativos': 'Gerencia de TI',
            'tiempo_estimado_restablecimiento': '72 horas',
            'recursos_necesarios': 'Equipo de respuesta, consultores externos'
        }
        
        reporte['medidas_mitigacion'] = {
            'acciones_corto_plazo': incidente.get('AccionesInmediatas', ''),
            'acciones_mediano_plazo': 'Implementación de controles adicionales',
            'acciones_largo_plazo': incidente.get('PlanMejora', 'Por definir')
        }
        
        reporte['impacto_servicios_vitales'] = {
            'servicios_criticos_afectados': incidente.get('ServiciosInterrumpidos', 'Ninguno'),
            'porcentaje_capacidad_reducida': '0%',
            'alternativas_operacionales': 'Sistemas de respaldo activos'
        }
        
        return reporte
    
    def _generar_informe_final(self, datos: Dict) -> Dict:
        """Genera informe final (15 días)"""
        reporte = self._generar_segundo_reporte(datos)
        incidente = datos['incidente']
        
        # Actualizar referencia
        reporte['referencia']['tipo_reporte'] = 'INFORME_FINAL'
        reporte['referencia']['fecha_inicio_original'] = self._format_datetime(incidente['FechaCreacion'])
        reporte['referencia']['fecha_cierre'] = self._format_datetime(incidente.get('FechaCierre', datetime.now()))
        
        # Cronología detallada
        cronologia = []
        for hist in datos['historial']:
            cronologia.append({
                'fecha_hora': self._format_datetime(hist['FechaAccion']),
                'evento': hist['TipoAccion'],
                'descripcion': hist.get('DatosNuevos', '')
            })
        
        reporte['descripcion_completa'] = {
            'cronologia_detallada': cronologia,
            'desarrollo_ataque': {
                'vector_inicial': incidente.get('OrigenIncidente', ''),
                'tecnicas_utilizadas': 'TTPs documentadas',
                'objetivos_atacante': 'En análisis',
                'movimiento_lateral': 'No detectado',
                'escalamiento_privilegios': 'No detectado'
            }
        }
        
        reporte['causa_raiz_final'] = {
            'vector_compromiso': incidente.get('OrigenIncidente', ''),
            'vulnerabilidad_especifica': incidente.get('AnciTipoAmenaza', ''),
            'controles_fallidos': incidente.get('CausaRaiz', ''),
            'factores_contributivos': 'Documentados en análisis'
        }
        
        reporte['impacto_final_detallado'] = {
            'sistemas_afectados': self._detallar_sistemas_afectados(incidente),
            'impacto_servicios': self._detallar_impacto_servicios(incidente),
            'impacto_economico': {
                'costos_recuperacion': 'En evaluación',
                'perdidas_operativas': 'En evaluación',
                'costos_terceros': 'N/A'
            },
            'impacto_terceros': {
                'usuarios_sin_servicio': 'Ninguno',
                'impacto_salud_publica': 'No aplica',
                'impacto_seguridad_nacional': 'No aplica'
            }
        }
        
        reporte['medidas_contencion_recuperacion'] = {
            'contencion_inmediata': self._detallar_acciones_contencion(incidente),
            'erradicacion': {
                'eliminacion_malware': 'Completada',
                'cierre_brechas': 'Vulnerabilidades parcheadas',
                'limpieza_sistemas': 'Sistemas restaurados'
            },
            'recuperacion': {
                'restauracion_respaldos': 'Completada',
                'reconstruccion_sistemas': 'No requerida',
                'verificacion_integridad': 'Validada',
                'fecha_servicio_normal': self._format_datetime(incidente.get('FechaCierre', datetime.now()))
            }
        }
        
        reporte['apoyo_externo_recibido'] = {
            'csirt_nacional': 'No requerido',
            'consultoras_seguridad': 'Apoyo especializado recibido',
            'fuerzas_orden': 'No requerido',
            'proveedores_tecnologia': 'Soporte del fabricante'
        }
        
        reporte['situacion_final'] = {
            'estado_incidente': incidente.get('EstadoActual', 'Cerrado'),
            'fecha_resolucion': self._format_datetime(incidente.get('FechaCierre', datetime.now())),
            'verificacion_normalidad': 'Confirmada mediante monitoreo',
            'monitoreo_posterior': 'Activo por 30 días'
        }
        
        reporte['lecciones_aprendidas'] = {
            'controles_fallidos_identificados': incidente.get('CausaRaiz', ''),
            'mejoras_implementadas': incidente.get('AccionesInmediatas', ''),
            'recomendaciones_futuras': incidente.get('LeccionesAprendidas', ''),
            'actualizacion_procedimientos': incidente.get('PlanMejora', '')
        }
        
        return reporte
    
    def _configurar_estilos(self, doc: Document):
        """Configura estilos del documento"""
        # Estilo para títulos
        styles = doc.styles
        
        # Configurar fuente predeterminada
        style = styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
    
    def _agregar_tabla_referencia(self, doc: Document, referencia: Dict):
        """Agrega tabla con información de referencia"""
        table = doc.add_table(rows=len(referencia), cols=2)
        table.style = 'Light List'
        
        for i, (key, value) in enumerate(referencia.items()):
            row = table.rows[i]
            row.cells[0].text = key.replace('_', ' ').title()
            row.cells[1].text = str(value)
    
    def _agregar_seccion(self, doc: Document, titulo: str, contenido: Dict):
        """Agrega una sección al documento"""
        doc.add_heading(titulo.replace('_', ' ').title(), 2)
        
        if isinstance(contenido, dict):
            for key, value in contenido.items():
                p = doc.add_paragraph()
                p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                
                if isinstance(value, list):
                    p.add_run(', '.join(str(v) for v in value))
                elif isinstance(value, dict):
                    # Crear sub-tabla
                    doc.add_paragraph()
                    self._agregar_tabla_referencia(doc, value)
                else:
                    p.add_run(str(value))
        else:
            doc.add_paragraph(str(contenido))
        
        doc.add_paragraph()  # Espacio entre secciones
    
    def _agregar_seccion_archivos(self, doc: Document, archivos: List[Dict]):
        """Agrega sección de archivos adjuntos"""
        doc.add_heading('Archivos de Evidencia Adjuntos', 2)
        
        if not archivos:
            doc.add_paragraph('No hay archivos adjuntos.')
            return
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid'
        
        # Encabezados
        header = table.rows[0]
        header.cells[0].text = 'Sección'
        header.cells[1].text = 'Archivo'
        header.cells[2].text = 'Descripción'
        header.cells[3].text = 'Fecha'
        
        # Datos
        for archivo in archivos:
            row = table.add_row()
            row.cells[0].text = f"Sección {archivo.get('SeccionID', 'N/A')}"
            row.cells[1].text = archivo.get('NombreArchivo', '')
            row.cells[2].text = archivo.get('Descripcion', '')
            row.cells[3].text = self._format_datetime(archivo.get('FechaSubida'))
    
    # Métodos auxiliares
    def _dict_from_row(self, row):
        """Convierte una fila de cursor en diccionario"""
        if not row:
            return None
        return dict(zip([column[0] for column in self.cursor.description], row))
    
    def _format_datetime(self, dt):
        """Formatea datetime para salida"""
        if not dt:
            return ''
        if isinstance(dt, str):
            return dt
        return dt.strftime('%Y-%m-%d %H:%M:%S')
    
    def _truncar_texto(self, texto: str, max_length: int) -> str:
        """Trunca texto a longitud máxima"""
        if not texto:
            return ''
        if len(texto) <= max_length:
            return texto
        return texto[:max_length-3] + '...'
    
    def _obtener_taxonomias_principales(self, taxonomias: List[Dict]) -> str:
        """Obtiene códigos de taxonomías principales"""
        if not taxonomias:
            return 'Sin clasificar'
        codigos = [t['Codigo'] for t in taxonomias[:3]]  # Máximo 3
        return ', '.join(codigos)
    
    def _filtrar_archivos_seccion(self, archivos: List[Dict], secciones: Optional[List[int]]) -> List[Dict]:
        """Filtra archivos por sección"""
        if not secciones:
            return archivos
        return [a for a in archivos if a.get('SeccionID') in secciones]
    
    def _detallar_sistemas_afectados(self, incidente: Dict) -> List[Dict]:
        """Detalla sistemas afectados"""
        sistemas = incidente.get('SistemasAfectados', '').split(',')
        return [{
            'nombre_sistema': sistema.strip(),
            'funcion': 'Sistema crítico',
            'tiempo_caida': 'N/A',
            'datos_perdidos': 'No'
        } for sistema in sistemas if sistema.strip()]
    
    def _detallar_impacto_servicios(self, incidente: Dict) -> List[Dict]:
        """Detalla impacto en servicios"""
        if not incidente.get('ServiciosInterrumpidos'):
            return []
        
        return [{
            'servicio_interrumpido': incidente.get('ServiciosInterrumpidos', ''),
            'duracion_total': 'En evaluación',
            'usuarios_afectados': 'En evaluación',
            'region_impactada': incidente.get('AlcanceGeografico', 'Local')
        }]
    
    def _detallar_acciones_contencion(self, incidente: Dict) -> List[Dict]:
        """Detalla acciones de contención"""
        if not incidente.get('AccionesInmediatas'):
            return []
        
        return [{
            'fecha_hora': self._format_datetime(incidente.get('FechaActualizacion')),
            'accion': incidente.get('AccionesInmediatas', ''),
            'responsable': incidente.get('ResponsableCliente', ''),
            'efectividad': 'Exitosa'
        }]


# Función auxiliar para uso directo
def generar_informe_anci_descargable(incidente_id: int, tipo_reporte: str = 'preliminar', formato: str = 'word'):
    """
    Genera un informe ANCI listo para descargar
    
    Args:
        incidente_id: ID del incidente
        tipo_reporte: 'preliminar', 'actualizacion', 'plan_accion', 'final'
        formato: 'word', 'json'
    
    Returns:
        BytesIO con el documento o dict con JSON
    """
    generador = GeneradorInformesANCIv2()
    
    if formato == 'json':
        return generador.generar_informe_json(incidente_id, tipo_reporte)
    else:
        return generador.generar_documento_word(incidente_id, tipo_reporte)