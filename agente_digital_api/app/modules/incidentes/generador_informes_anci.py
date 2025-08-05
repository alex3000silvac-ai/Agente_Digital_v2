#!/usr/bin/env python3
"""
Generador de Informes ANCI
Genera reportes de incidentes basados en plantillas Word (.docx)
"""

import os
import json
from datetime import datetime
from typing import Dict, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..core.database import get_db_connection
from ..incidentes.sistema_dinamico import SistemaDinamicoIncidentes

class GeneradorInformesANCI:
    """
    Genera informes ANCI basados en plantillas Word
    """
    
    def __init__(self):
        self.plantilla_path = os.environ.get('ANCI_TEMPLATE_PATH', 
            'C:/Users/alexs/Downloads/Informes ANCI_ reporte de incidentes_.docx')
        self.sistema = SistemaDinamicoIncidentes()
        
    def generar_informe(self, incidente_id: int, tipo_informe: str = 'completo') -> str:
        """
        Genera un informe ANCI para un incidente
        
        Args:
            incidente_id: ID del incidente
            tipo_informe: 'completo', 'preliminar', 'final'
            
        Returns:
            Ruta del archivo generado
        """
        try:
            # Verificar que existe la plantilla
            if not os.path.exists(self.plantilla_path):
                # Buscar en rutas alternativas
                rutas_alternativas = [
                    './plantillas/ANCI_template.docx',
                    '../plantillas/ANCI_template.docx',
                    '/plantillas/ANCI_template.docx',
                    'C:/plantillas/ANCI_template.docx'
                ]
                
                for ruta in rutas_alternativas:
                    if os.path.exists(ruta):
                        self.plantilla_path = ruta
                        break
                else:
                    raise FileNotFoundError(f"No se encontró la plantilla ANCI en ninguna ubicación")
            
            # Cargar datos del incidente
            datos_incidente = self._cargar_datos_incidente(incidente_id)
            
            # Cargar plantilla
            doc = Document(self.plantilla_path)
            
            # Reemplazar marcadores en el documento
            self._reemplazar_marcadores(doc, datos_incidente, tipo_informe)
            
            # Agregar secciones dinámicas
            self._agregar_secciones_dinamicas(doc, datos_incidente)
            
            # Agregar evidencias como anexos
            self._agregar_anexos_evidencias(doc, datos_incidente)
            
            # Guardar documento
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            nombre_archivo = f"Informe_ANCI_{datos_incidente['IDVisible']}_{tipo_informe}_{timestamp}.docx"
            
            # Crear carpeta de informes si no existe
            carpeta_informes = os.path.join(
                os.environ.get('ARCHIVOS_PATH', '/archivos'),
                f"empresa_{datos_incidente['EmpresaID']}",
                f"incidente_{datos_incidente['IDVisible']}",
                'informes_anci'
            )
            os.makedirs(carpeta_informes, exist_ok=True)
            
            ruta_completa = os.path.join(carpeta_informes, nombre_archivo)
            doc.save(ruta_completa)
            
            # Registrar en auditoría
            self._registrar_generacion_informe(incidente_id, tipo_informe, ruta_completa)
            
            return ruta_completa
            
        except Exception as e:
            raise Exception(f"Error generando informe ANCI: {e}")
    
    def _cargar_datos_incidente(self, incidente_id: int) -> Dict:
        """Carga todos los datos del incidente para el informe"""
        # Usar el sistema dinámico para cargar todo
        datos_completos = self.sistema.cargar_incidente_completo(incidente_id)
        
        if not datos_completos['success']:
            raise ValueError(f"No se pudo cargar el incidente {incidente_id}")
        
        return datos_completos
    
    def _reemplazar_marcadores(self, doc: Document, datos: Dict, tipo_informe: str):
        """Reemplaza los marcadores en el documento con datos reales"""
        
        incidente = datos['incidente']
        
        # Marcadores básicos a reemplazar
        marcadores = {
            '{{FECHA_REPORTE}}': datetime.now().strftime('%d/%m/%Y'),
            '{{TIPO_REPORTE}}': tipo_informe.upper(),
            '{{ID_INCIDENTE}}': str(incidente.get('IDVisible', '')),
            '{{TITULO_INCIDENTE}}': incidente.get('Titulo', ''),
            '{{FECHA_DETECCION}}': self._formatear_fecha(incidente.get('FechaDeteccion')),
            '{{FECHA_OCURRENCIA}}': self._formatear_fecha(incidente.get('FechaOcurrencia')),
            '{{CRITICIDAD}}': incidente.get('Criticidad', ''),
            '{{ESTADO}}': incidente.get('EstadoActual', ''),
            '{{EMPRESA_ID}}': str(incidente.get('EmpresaID', '')),
            '{{TIPO_EMPRESA}}': incidente.get('Tipo_Empresa', ''),
            '{{ORIGEN_INCIDENTE}}': incidente.get('OrigenIncidente', ''),
            '{{SISTEMAS_AFECTADOS}}': incidente.get('SistemasAfectados', ''),
            '{{SERVICIOS_INTERRUMPIDOS}}': incidente.get('ServiciosInterrumpidos', ''),
            '{{ALCANCE_GEOGRAFICO}}': incidente.get('AlcanceGeografico', ''),
            '{{RESPONSABLE_CLIENT}}': incidente.get('ResponsableCliente', ''),
            '{{REPORTE_ANCI_ID}}': incidente.get('ReporteAnciID', ''),
            '{{FECHA_DECLARACION_ANCI}}': self._formatear_fecha(incidente.get('FechaDeclaracionANCI')),
            '{{TIPO_AMENAZA}}': incidente.get('AnciTipoAmenaza', ''),
            '{{IMPACTO_PRELIMINAR}}': incidente.get('AnciImpactoPreliminar', ''),
            '{{CAUSA_RAIZ}}': incidente.get('CausaRaiz', ''),
            '{{LECCIONES_APRENDIDAS}}': incidente.get('LeccionesAprendidas', ''),
            '{{PLAN_MEJORA}}': incidente.get('PlanMejora', '')
        }
        
        # Reemplazar en párrafos
        for paragraph in doc.paragraphs:
            for marcador, valor in marcadores.items():
                if marcador in paragraph.text:
                    paragraph.text = paragraph.text.replace(marcador, valor or 'N/A')
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for marcador, valor in marcadores.items():
                        if marcador in cell.text:
                            cell.text = cell.text.replace(marcador, valor or 'N/A')
    
    def _agregar_secciones_dinamicas(self, doc: Document, datos: Dict):
        """Agrega las secciones dinámicas según el tipo de empresa"""
        
        # Agregar título de secciones
        doc.add_page_break()
        titulo = doc.add_heading('DETALLE DE SECCIONES', level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Iterar por cada sección con contenido
        for seccion in datos['secciones']:
            if seccion['tiene_contenido']:
                # Título de sección
                doc.add_heading(f"{seccion['titulo']}", level=2)
                
                # Descripción
                if seccion['descripcion']:
                    p = doc.add_paragraph(seccion['descripcion'])
                    p.style = 'Normal'
                
                # Datos de la sección
                if seccion['datos']:
                    doc.add_heading('Información:', level=3)
                    for campo, valor in seccion['datos'].items():
                        if valor:
                            p = doc.add_paragraph()
                            p.add_run(f"{self._formatear_nombre_campo(campo)}: ").bold = True
                            p.add_run(str(valor))
                
                # Comentarios
                if seccion['comentarios']:
                    doc.add_heading('Comentarios:', level=3)
                    for comentario in seccion['comentarios']:
                        p = doc.add_paragraph(f"• {comentario['texto']}")
                        p.add_run(f" ({comentario['usuario']} - {self._formatear_fecha(comentario['fecha'])})").font.size = Pt(9)
                
                # Archivos adjuntos
                if seccion['archivos']:
                    doc.add_heading('Evidencias adjuntas:', level=3)
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    
                    # Headers
                    headers = ['#', 'Archivo', 'Tamaño', 'Fecha']
                    for i, header in enumerate(headers):
                        table.rows[0].cells[i].text = header
                    
                    # Datos de archivos
                    for archivo in seccion['archivos']:
                        row = table.add_row()
                        row.cells[0].text = str(archivo['numero'])
                        row.cells[1].text = archivo['nombre_original']
                        row.cells[2].text = f"{archivo['tamano_kb']} KB"
                        row.cells[3].text = self._formatear_fecha(archivo['fecha'])
                
                doc.add_paragraph()  # Espacio entre secciones
    
    def _agregar_anexos_evidencias(self, doc: Document, datos: Dict):
        """Agrega lista de todas las evidencias como anexo"""
        
        doc.add_page_break()
        doc.add_heading('ANEXO: LISTADO COMPLETO DE EVIDENCIAS', level=1)
        
        total_archivos = sum(len(s['archivos']) for s in datos['secciones'])
        
        if total_archivos == 0:
            doc.add_paragraph('No hay evidencias adjuntas a este incidente.')
            return
        
        # Crear tabla resumen
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['Sección', 'Tipo', 'Archivo', 'Descripción', 'Fecha']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # Estilo para headers
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Agregar todas las evidencias
        for seccion in datos['secciones']:
            for archivo in seccion['archivos']:
                row = table.add_row()
                row.cells[0].text = seccion['titulo']
                row.cells[1].text = seccion['tipo']
                row.cells[2].text = archivo['nombre_original']
                row.cells[3].text = archivo.get('descripcion', '')
                row.cells[4].text = self._formatear_fecha(archivo['fecha'])
        
        # Nota al pie
        doc.add_paragraph()
        nota = doc.add_paragraph()
        nota.add_run('Nota: ').bold = True
        nota.add_run('Los archivos físicos se encuentran almacenados en el sistema y están disponibles para consulta.')
    
    def _formatear_fecha(self, fecha) -> str:
        """Formatea una fecha para el informe"""
        if not fecha:
            return 'N/A'
        
        try:
            if isinstance(fecha, str):
                fecha_obj = datetime.fromisoformat(fecha.replace('Z', '+00:00'))
            else:
                fecha_obj = fecha
            
            return fecha_obj.strftime('%d/%m/%Y %H:%M')
        except:
            return str(fecha)
    
    def _formatear_nombre_campo(self, campo: str) -> str:
        """Formatea el nombre de un campo para mostrar"""
        return campo.replace('_', ' ').title()
    
    def _registrar_generacion_informe(self, incidente_id: int, tipo_informe: str, ruta_archivo: str):
        """Registra en auditoría la generación del informe"""
        conn = None
        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Si existe la tabla de auditoría
            cursor.execute("""
                IF EXISTS (SELECT * FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'INCIDENTES_AUDITORIA')
                BEGIN
                    INSERT INTO INCIDENTES_AUDITORIA 
                    (IncidenteID, TipoAccion, DatosNuevos, Usuario, FechaAccion)
                    VALUES (?, ?, ?, ?, GETDATE())
                END
            """, (
                incidente_id,
                'GENERAR_INFORME_ANCI',
                json.dumps({
                    'tipo_informe': tipo_informe,
                    'ruta_archivo': ruta_archivo,
                    'plantilla': self.plantilla_path
                }),
                'Sistema'
            ))
            
            conn.commit()
            
        except Exception as e:
            print(f"Error registrando auditoría: {e}")
        finally:
            if conn:
                conn.close()
    
    def listar_plantillas_disponibles(self) -> list:
        """Lista todas las plantillas ANCI disponibles"""
        plantillas = []
        
        # Buscar en varias ubicaciones
        ubicaciones = [
            'C:/Users/alexs/Downloads/',
            './plantillas/',
            '../plantillas/',
            '/plantillas/',
            os.environ.get('ANCI_TEMPLATES_DIR', './plantillas')
        ]
        
        for ubicacion in ubicaciones:
            if os.path.exists(ubicacion):
                for archivo in os.listdir(ubicacion):
                    if archivo.endswith('.docx') and 'ANCI' in archivo.upper():
                        plantillas.append({
                            'nombre': archivo,
                            'ruta': os.path.join(ubicacion, archivo),
                            'tamano': os.path.getsize(os.path.join(ubicacion, archivo))
                        })
        
        return plantillas