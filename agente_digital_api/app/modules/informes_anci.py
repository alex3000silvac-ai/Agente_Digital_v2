"""
Módulo para generación de informes ANCI
Maneja la lógica de generación de informes preliminares, completos y finales
"""

import os
import json
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.database import get_db_connection
from config import Config
UPLOAD_FOLDER = Config.UPLOAD_FOLDER

logger = logging.getLogger(__name__)

class InformesANCI:
    def __init__(self):
        self.templates_dir = os.path.join(os.path.dirname(__file__), 'templates_anci')
        self.output_dir = os.path.join(UPLOAD_FOLDER, 'informes_anci')
        
        # Crear directorios si no existen
        os.makedirs(self.templates_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Marcadores por tipo de informe
        self.marcadores_informe = {
            'preliminar': [
                '{{FECHA_REPORTE}}',
                '{{ID_INCIDENTE}}',
                '{{EMPRESA_NOMBRE}}',
                '{{EMPRESA_RUT}}',
                '{{TITULO_INCIDENTE}}',
                '{{FECHA_DETECCION}}',
                '{{DESCRIPCION_BREVE}}',
                '{{IMPACTO_INICIAL}}',
                '{{ACCIONES_INMEDIATAS}}',
                '{{CONTACTO_RESPONSABLE}}'
            ],
            'completo': [
                # Incluye todos los del preliminar más:
                '{{ANALISIS_DETALLADO}}',
                '{{SISTEMAS_AFECTADOS}}',
                '{{TAXONOMIAS_SELECCIONADAS}}',
                '{{EVIDENCIAS_RECOPILADAS}}',
                '{{CRONOLOGIA_EVENTOS}}',
                '{{MEDIDAS_CONTENCION}}',
                '{{PLAN_RECUPERACION}}'
            ],
            'final': [
                # Incluye todos los anteriores más:
                '{{CAUSA_RAIZ}}',
                '{{LECCIONES_APRENDIDAS}}',
                '{{MEJORAS_IMPLEMENTADAS}}',
                '{{RECOMENDACIONES}}',
                '{{COSTO_TOTAL}}',
                '{{TIEMPO_RESOLUCION}}',
                '{{METRICAS_FINALES}}'
            ]
        }
    
    def generar_informe(self, incidente_id, tipo_informe='preliminar', usuario_id='sistema'):
        """
        Genera un informe ANCI del tipo especificado
        """
        try:
            logger.info(f"Iniciando generación de informe {tipo_informe} para incidente {incidente_id}")
            
            # Validar tipo de informe
            if tipo_informe not in ['preliminar', 'completo', 'final']:
                return {
                    'exito': False,
                    'error': 'Tipo de informe no válido'
                }
            
            # Obtener datos del incidente
            datos_incidente = self._obtener_datos_incidente(incidente_id)
            if not datos_incidente:
                return {
                    'exito': False,
                    'error': 'Incidente no encontrado'
                }
            
            # Cargar archivo semilla si existe
            datos_semilla = self._cargar_semilla(incidente_id)
            
            # Combinar datos
            datos_completos = {**datos_incidente, **datos_semilla}
            
            # Generar el documento
            if tipo_informe == 'preliminar':
                documento = self._generar_informe_preliminar(datos_completos)
            elif tipo_informe == 'completo':
                documento = self._generar_informe_completo(datos_completos)
            else:  # final
                documento = self._generar_informe_final(datos_completos)
            
            # Guardar el documento
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            nombre_archivo = f"ANCI_{tipo_informe}_{datos_incidente['IDVisible']}_{timestamp}.docx"
            ruta_archivo = os.path.join(self.output_dir, nombre_archivo)
            
            documento.save(ruta_archivo)
            
            # Registrar en base de datos
            informe_id = self._registrar_informe(
                incidente_id=incidente_id,
                tipo_informe=tipo_informe,
                ruta_archivo=ruta_archivo,
                usuario_id=usuario_id
            )
            
            logger.info(f"Informe generado exitosamente: {ruta_archivo}")
            
            return {
                'exito': True,
                'informe_id': informe_id,
                'ruta_archivo': ruta_archivo
            }
            
        except Exception as e:
            logger.error(f"Error generando informe: {str(e)}")
            return {
                'exito': False,
                'error': str(e)
            }
    
    def _obtener_datos_incidente(self, incidente_id):
        """
        Obtiene los datos básicos del incidente desde la BD
        """
        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            
            query = """
            SELECT 
                i.IncidenteID,
                i.IDVisible,
                i.Titulo,
                i.FechaDeteccion,
                i.FechaOcurrencia,
                i.OrigenIncidente,
                i.Criticidad,
                i.EstadoActual,
                i.DescripcionBreve,
                i.ImpactoPreliminar,
                i.AccionesInmediatas,
                i.TipoRegistro,
                i.SolicitarCSIRT,
                e.Nombre as EmpresaNombre,
                e.RUT as EmpresaRUT,
                e.TipoEmpresa,
                inq.Nombre as InquilinoNombre
            FROM Incidentes i
            INNER JOIN Empresas e ON i.EmpresaID = e.ID
            INNER JOIN Inquilinos inq ON e.InquilinoID = inq.ID
            WHERE i.IncidenteID = ? AND i.Activo = 1
            """
            
            cursor.execute(query, (incidente_id,))
            columns = [column[0] for column in cursor.description]
            row = cursor.fetchone()
            
            if not row:
                return None
                
            datos = dict(zip(columns, row))
            
            # Convertir fechas a string
            for campo in ['FechaDeteccion', 'FechaOcurrencia']:
                if datos.get(campo):
                    datos[campo] = datos[campo].strftime('%d/%m/%Y %H:%M')
            
            # Obtener taxonomías seleccionadas
            datos['Taxonomias'] = self._obtener_taxonomias_incidente(cursor, incidente_id)
            
            # Obtener archivos adjuntos
            datos['Archivos'] = self._obtener_archivos_incidente(cursor, incidente_id)
            
            return datos
            
        except Exception as e:
            logger.error(f"Error obteniendo datos del incidente: {str(e)}")
            return None
        finally:
            if conn:
                conn.close()
    
    def _obtener_taxonomias_incidente(self, cursor, incidente_id):
        """
        Obtiene las taxonomías seleccionadas para el incidente
        """
        query = """
        SELECT 
            t.Area,
            t.Categoria_del_Incidente,
            t.Subcategoria_del_Incidente,
            t.Efecto,
            it.Justificacion
        FROM INCIDENTE_TAXONOMIA it
        INNER JOIN TAXONOMIA_INCIDENTES t ON it.Id_Taxonomia = t.Id_Incidente
        WHERE it.IncidenteID = ?
        ORDER BY t.Area, t.Categoria_del_Incidente
        """
        
        cursor.execute(query, (incidente_id,))
        taxonomias = []
        
        for row in cursor.fetchall():
            taxonomias.append({
                'area': row[0],
                'categoria': row[1],
                'subcategoria': row[2],
                'efecto': row[3],
                'justificacion': row[4]
            })
            
        return taxonomias
    
    def _obtener_archivos_incidente(self, cursor, incidente_id):
        """
        Obtiene información de archivos adjuntos
        """
        query = """
        SELECT 
            SeccionFormulario,
            NombreArchivo,
            TamanoKB,
            FechaSubida
        FROM EVIDENCIAS_INCIDENTE
        WHERE IncidenteID = ?
        ORDER BY SeccionFormulario, NumeroEvidencia
        """
        
        cursor.execute(query, (incidente_id,))
        archivos = []
        
        for row in cursor.fetchall():
            archivos.append({
                'seccion': row[0],
                'nombre': row[1],
                'tamano_kb': row[2],
                'fecha': row[3].strftime('%d/%m/%Y %H:%M') if row[3] else ''
            })
            
        return archivos
    
    def _cargar_semilla(self, incidente_id):
        """
        Carga el archivo semilla JSON del incidente si existe
        """
        try:
            ruta_semilla = os.path.join(
                UPLOAD_FOLDER,
                'archivos_temporales',
                f'incidente_{incidente_id}_datos.json'
            )
            
            if os.path.exists(ruta_semilla):
                with open(ruta_semilla, 'r', encoding='utf-8') as f:
                    return json.load(f)
            else:
                logger.warning(f"Archivo semilla no encontrado: {ruta_semilla}")
                return {}
                
        except Exception as e:
            logger.error(f"Error cargando semilla: {str(e)}")
            return {}
    
    def _generar_informe_preliminar(self, datos):
        """
        Genera un informe preliminar (24 horas)
        """
        doc = Document()
        
        # Título principal
        titulo = doc.add_heading('INFORME PRELIMINAR ANCI', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo
        subtitulo = doc.add_heading('Notificación de Incidente de Ciberseguridad', level=2)
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del reporte
        doc.add_paragraph()
        info = doc.add_table(rows=3, cols=2)
        info.style = 'Light Grid Accent 1'
        
        info.cell(0, 0).text = 'Fecha del Reporte:'
        info.cell(0, 1).text = datetime.now().strftime('%d/%m/%Y %H:%M')
        info.cell(1, 0).text = 'ID Incidente:'
        info.cell(1, 1).text = datos.get('IDVisible', 'N/A')
        info.cell(2, 0).text = 'Tipo de Informe:'
        info.cell(2, 1).text = 'PRELIMINAR (24 horas)'
        
        # Sección 1: Información de la Empresa
        doc.add_heading('1. INFORMACIÓN DE LA EMPRESA', level=1)
        doc.add_paragraph(f"Nombre: {datos.get('EmpresaNombre', 'N/A')}")
        doc.add_paragraph(f"RUT: {datos.get('EmpresaRUT', 'N/A')}")
        doc.add_paragraph(f"Tipo: {datos.get('TipoEmpresa', 'N/A')}")
        doc.add_paragraph(f"Inquilino: {datos.get('InquilinoNombre', 'N/A')}")
        
        # Sección 2: Información del Incidente
        doc.add_heading('2. INFORMACIÓN DEL INCIDENTE', level=1)
        doc.add_paragraph(f"Título: {datos.get('Titulo', 'N/A')}")
        doc.add_paragraph(f"Fecha de Detección: {datos.get('FechaDeteccion', 'N/A')}")
        doc.add_paragraph(f"Fecha de Ocurrencia: {datos.get('FechaOcurrencia', 'N/A')}")
        doc.add_paragraph(f"Criticidad: {datos.get('Criticidad', 'N/A')}")
        doc.add_paragraph(f"Origen: {datos.get('OrigenIncidente', 'N/A')}")
        
        # Sección 3: Descripción Preliminar
        doc.add_heading('3. DESCRIPCIÓN PRELIMINAR', level=1)
        doc.add_paragraph(datos.get('DescripcionBreve', 'Sin descripción disponible'))
        
        # Sección 4: Impacto Inicial
        doc.add_heading('4. EVALUACIÓN DE IMPACTO INICIAL', level=1)
        doc.add_paragraph(datos.get('ImpactoPreliminar', 'Evaluación de impacto pendiente'))
        
        # Sección 5: Acciones Inmediatas
        doc.add_heading('5. ACCIONES INMEDIATAS TOMADAS', level=1)
        doc.add_paragraph(datos.get('AccionesInmediatas', 'Sin acciones registradas'))
        
        # Sección 6: Próximos Pasos
        doc.add_heading('6. PRÓXIMOS PASOS', level=1)
        doc.add_paragraph('• Continuar con la investigación del incidente')
        doc.add_paragraph('• Recopilar evidencia adicional')
        doc.add_paragraph('• Evaluar el alcance completo del impacto')
        doc.add_paragraph('• Preparar informe completo en las próximas 72 horas')
        
        # Nota al pie
        doc.add_paragraph()
        doc.add_paragraph(
            'Este es un informe preliminar generado dentro de las 24 horas posteriores '
            'a la detección del incidente, según lo requerido por la normativa ANCI.',
            style='Quote'
        )
        
        return doc
    
    def _generar_informe_completo(self, datos):
        """
        Genera un informe completo (72 horas)
        """
        # Empezar con el contenido del preliminar
        doc = self._generar_informe_preliminar(datos)
        
        # Agregar secciones adicionales
        doc.add_page_break()
        doc.add_heading('INFORMACIÓN ADICIONAL - INFORME COMPLETO', level=1)
        
        # Sección 7: Análisis Detallado
        doc.add_heading('7. ANÁLISIS DETALLADO', level=1)
        
        # Sistemas afectados
        doc.add_heading('7.1 Sistemas Afectados', level=2)
        sistemas = datos.get('sistemas_afectados', {})
        if sistemas:
            for sistema, detalle in sistemas.items():
                doc.add_paragraph(f"• {sistema}: {detalle}")
        else:
            doc.add_paragraph('Información de sistemas afectados pendiente de análisis')
        
        # Taxonomías
        doc.add_heading('7.2 Clasificación según Taxonomía ANCI', level=2)
        if datos.get('Taxonomias'):
            tabla_tax = doc.add_table(rows=1, cols=4)
            tabla_tax.style = 'Light Grid Accent 1'
            
            # Headers
            headers = ['Área', 'Categoría', 'Efecto', 'Justificación']
            for i, header in enumerate(headers):
                tabla_tax.cell(0, i).text = header
            
            # Datos
            for tax in datos['Taxonomias']:
                row = tabla_tax.add_row()
                row.cells[0].text = tax.get('area', '')
                row.cells[1].text = tax.get('categoria', '')
                row.cells[2].text = tax.get('efecto', '')
                row.cells[3].text = tax.get('justificacion', '')
        else:
            doc.add_paragraph('No se han seleccionado taxonomías para este incidente')
        
        # Sección 8: Evidencias
        doc.add_heading('8. EVIDENCIAS RECOPILADAS', level=1)
        if datos.get('Archivos'):
            for i, archivo in enumerate(datos['Archivos'], 1):
                doc.add_paragraph(
                    f"{i}. {archivo['nombre']} - Sección: {archivo['seccion']} "
                    f"({archivo['tamano_kb']} KB) - {archivo['fecha']}"
                )
        else:
            doc.add_paragraph('No se han adjuntado evidencias')
        
        # Sección 9: Plan de Recuperación
        doc.add_heading('9. PLAN DE RECUPERACIÓN', level=1)
        plan = datos.get('plan_recuperacion', {})
        if plan:
            for fase, actividades in plan.items():
                doc.add_paragraph(f"{fase}:")
                for actividad in actividades:
                    doc.add_paragraph(f"  • {actividad}")
        else:
            doc.add_paragraph('Plan de recuperación en desarrollo')
        
        return doc
    
    def _generar_informe_final(self, datos):
        """
        Genera un informe final (30 días)
        """
        # Empezar con el contenido del completo
        doc = self._generar_informe_completo(datos)
        
        # Agregar secciones finales
        doc.add_page_break()
        doc.add_heading('ANÁLISIS FINAL Y CONCLUSIONES', level=1)
        
        # Sección 10: Causa Raíz
        doc.add_heading('10. ANÁLISIS DE CAUSA RAÍZ', level=1)
        causa_raiz = datos.get('causa_raiz', 'Análisis de causa raíz pendiente')
        doc.add_paragraph(causa_raiz)
        
        # Sección 11: Lecciones Aprendidas
        doc.add_heading('11. LECCIONES APRENDIDAS', level=1)
        lecciones = datos.get('lecciones_aprendidas', [])
        if lecciones:
            for leccion in lecciones:
                doc.add_paragraph(f"• {leccion}")
        else:
            doc.add_paragraph('Documentación de lecciones aprendidas en proceso')
        
        # Sección 12: Mejoras Implementadas
        doc.add_heading('12. MEJORAS IMPLEMENTADAS', level=1)
        mejoras = datos.get('mejoras_implementadas', [])
        if mejoras:
            for mejora in mejoras:
                doc.add_paragraph(f"• {mejora}")
        else:
            doc.add_paragraph('Documentación de mejoras en proceso')
        
        # Sección 13: Métricas Finales
        doc.add_heading('13. MÉTRICAS DEL INCIDENTE', level=1)
        doc.add_paragraph(f"Tiempo total de resolución: {datos.get('tiempo_resolucion', 'N/A')}")
        doc.add_paragraph(f"Costo estimado: {datos.get('costo_total', 'N/A')}")
        doc.add_paragraph(f"Usuarios afectados: {datos.get('usuarios_afectados', 'N/A')}")
        doc.add_paragraph(f"Servicios impactados: {datos.get('servicios_impactados', 'N/A')}")
        
        # Sección 14: Recomendaciones
        doc.add_heading('14. RECOMENDACIONES', level=1)
        recomendaciones = datos.get('recomendaciones', [])
        if recomendaciones:
            for rec in recomendaciones:
                doc.add_paragraph(f"• {rec}")
        else:
            doc.add_paragraph('• Implementar monitoreo continuo')
            doc.add_paragraph('• Actualizar procedimientos de respuesta')
            doc.add_paragraph('• Realizar capacitación al personal')
            doc.add_paragraph('• Revisar y actualizar políticas de seguridad')
        
        # Conclusión
        doc.add_heading('15. CONCLUSIÓN', level=1)
        doc.add_paragraph(
            'Este informe final documenta el ciclo completo de gestión del incidente, '
            'desde su detección hasta su resolución completa, incluyendo las lecciones '
            'aprendidas y mejoras implementadas para prevenir incidentes similares en el futuro.'
        )
        
        return doc
    
    def _registrar_informe(self, incidente_id, tipo_informe, ruta_archivo, usuario_id):
        """
        Registra el informe generado en la base de datos
        """
        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Obtener tamaño del archivo
            tamano_kb = os.path.getsize(ruta_archivo) / 1024
            
            # Insertar registro
            query = """
            INSERT INTO INFORMES_ANCI (
                IncidenteID,
                TipoInforme,
                EstadoInforme,
                FechaGeneracion,
                RutaArchivo,
                TamanoKB,
                GeneradoPor,
                Version,
                Activo
            ) VALUES (?, ?, ?, GETDATE(), ?, ?, ?, 
                (SELECT ISNULL(MAX(Version), 0) + 1 
                 FROM INFORMES_ANCI 
                 WHERE IncidenteID = ? AND TipoInforme = ?),
                1)
            """
            
            cursor.execute(query, (
                incidente_id,
                tipo_informe,
                'generado',
                ruta_archivo,
                tamano_kb,
                usuario_id,
                incidente_id,
                tipo_informe
            ))
            
            conn.commit()
            informe_id = cursor.lastrowid
            
            logger.info(f"Informe registrado en BD con ID: {informe_id}")
            return informe_id
            
        except Exception as e:
            logger.error(f"Error registrando informe en BD: {str(e)}")
            raise
        finally:
            if conn:
                conn.close()
    
    def validar_datos_incidente(self, incidente_id):
        """
        Valida que el incidente tenga los datos mínimos para generar informes
        """
        try:
            datos = self._obtener_datos_incidente(incidente_id)
            
            if not datos:
                return {
                    'valido': False,
                    'errores': ['Incidente no encontrado']
                }
            
            errores = []
            advertencias = []
            
            # Validar campos obligatorios
            campos_obligatorios = [
                'Titulo', 'FechaDeteccion', 'DescripcionBreve',
                'EmpresaNombre', 'EmpresaRUT'
            ]
            
            for campo in campos_obligatorios:
                if not datos.get(campo):
                    errores.append(f"Campo obligatorio faltante: {campo}")
            
            # Validar campos recomendados
            campos_recomendados = [
                'ImpactoPreliminar', 'AccionesInmediatas',
                'OrigenIncidente', 'Criticidad'
            ]
            
            for campo in campos_recomendados:
                if not datos.get(campo):
                    advertencias.append(f"Campo recomendado faltante: {campo}")
            
            # Validar taxonomías
            if not datos.get('Taxonomias'):
                advertencias.append('No se han seleccionado taxonomías')
            
            # Validar archivos
            if not datos.get('Archivos'):
                advertencias.append('No se han adjuntado evidencias')
            
            return {
                'valido': len(errores) == 0,
                'errores': errores,
                'advertencias': advertencias,
                'datos_disponibles': {
                    'titulo': bool(datos.get('Titulo')),
                    'fecha_deteccion': bool(datos.get('FechaDeteccion')),
                    'descripcion': bool(datos.get('DescripcionBreve')),
                    'impacto': bool(datos.get('ImpactoPreliminar')),
                    'acciones': bool(datos.get('AccionesInmediatas')),
                    'taxonomias': len(datos.get('Taxonomias', [])),
                    'archivos': len(datos.get('Archivos', []))
                }
            }
            
        except Exception as e:
            logger.error(f"Error validando datos: {str(e)}")
            return {
                'valido': False,
                'errores': [str(e)]
            }