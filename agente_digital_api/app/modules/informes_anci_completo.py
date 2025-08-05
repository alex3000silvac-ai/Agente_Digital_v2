#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de Informes ANCI Completo
Incluye todos los campos requeridos por la normativa ANCI
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from ..database import get_db_connection
import os
import traceback

class GeneradorInformeANCICompleto:
    """Genera informes ANCI con todos los campos requeridos"""
    
    def __init__(self):
        self.conn = None
        self.cursor = None
    
    def generar_informe_docx(self, reporte_id):
        """
        Genera un informe ANCI completo en formato Word
        
        Args:
            reporte_id: ID del reporte ANCI
            
        Returns:
            Ruta del archivo generado
        """
        try:
            self.conn = get_db_connection()
            self.cursor = self.conn.cursor()
            
            # Obtener datos completos del reporte
            datos = self._obtener_datos_completos(reporte_id)
            
            if not datos:
                raise ValueError(f"No se encontró el reporte ANCI {reporte_id}")
            
            # Crear documento
            doc = Document()
            self._configurar_estilos(doc)
            
            # Generar contenido
            self._agregar_encabezado(doc, datos)
            self._agregar_seccion_1_reportante(doc, datos)
            self._agregar_seccion_2_identificacion(doc, datos)
            self._agregar_seccion_3_medidas(doc, datos)
            
            # Si es OIV, agregar sección 4
            if datos.get('TipoEmpresa') == 'OIV':
                self._agregar_seccion_4_oiv(doc, datos)
            
            self._agregar_anexos(doc, datos)
            self._agregar_contacto_seguimiento(doc, datos)
            
            # Guardar documento
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"Informe_ANCI_Completo_{datos.get('IncidenteID', 'X')}_{timestamp}.docx"
            
            # Crear directorio si no existe
            carpeta = os.path.join(
                os.path.dirname(__file__), 
                '..', '..', 
                'informes_anci'
            )
            os.makedirs(carpeta, exist_ok=True)
            
            filepath = os.path.join(carpeta, filename)
            doc.save(filepath)
            
            # Agregar comentario sobre la protección de la sección 1
            print(f"\n⚠️  IMPORTANTE: La Sección 1 del documento está marcada como protegida.")
            print(f"   No debe ser editada una vez generada.")
            
            return filepath
            
        except Exception as e:
            print(f"Error generando informe: {e}")
            traceback.print_exc()
            raise
        finally:
            if self.conn:
                self.conn.close()
    
    def _obtener_datos_completos(self, reporte_id):
        """Obtiene todos los datos del reporte ANCI incluyendo campos faltantes"""
        query = """
        SELECT 
            r.*,
            i.IncidenteID,
            i.Titulo,
            i.FechaDeteccion,
            i.FechaOcurrencia,
            i.Criticidad,
            i.TipoRegistro,
            i.OrigenIncidente,
            i.SistemasAfectados,
            i.ServiciosInterrumpidos,
            i.AlcanceGeografico,
            i.ResponsableCliente,
            i.DescripcionInicial,
            i.CausaRaiz,
            i.LeccionesAprendidas,
            i.PlanMejora,
            e.RazonSocial as EmpresaNombre,
            e.RUT as EmpresaRUT,
            e.TipoEmpresa
        FROM ReportesANCI r
        LEFT JOIN Incidentes i ON r.IncidenteID = i.IncidenteID
        LEFT JOIN Empresas e ON i.EmpresaID = e.EmpresaID
        WHERE r.ReporteAnciID = ?
        """
        
        self.cursor.execute(query, (reporte_id,))
        columns = [column[0] for column in self.cursor.description]
        row = self.cursor.fetchone()
        
        if not row:
            return None
        
        datos = dict(zip(columns, row))
        
        # Obtener taxonomías asociadas
        datos['taxonomias'] = self._obtener_taxonomias(datos.get('IncidenteID'))
        
        # Obtener evidencias
        datos['evidencias'] = self._obtener_evidencias(datos.get('IncidenteID'))
        
        return datos
    
    def _obtener_taxonomias(self, incidente_id):
        """Obtiene las taxonomías asociadas al incidente"""
        if not incidente_id:
            return []
        
        query = """
        SELECT 
            ti.Area,
            ti.Efecto,
            ti.Categoria_del_Incidente as Categoria,
            ti.Descripcion
        FROM INCIDENTE_TAXONOMIA it
        INNER JOIN Taxonomia_incidentes ti ON it.Id_Taxonomia = ti.Id_Incidente
        WHERE it.IncidenteID = ?
        """
        
        self.cursor.execute(query, (incidente_id,))
        return self.cursor.fetchall()
    
    def _obtener_evidencias(self, incidente_id):
        """Obtiene las evidencias del incidente"""
        if not incidente_id:
            return []
        
        query = """
        SELECT 
            SeccionID,
            NombreOriginal,
            TamanoKB,
            FechaSubida,
            SubidoPor
        FROM INCIDENTES_ARCHIVOS
        WHERE IncidenteID = ?
        ORDER BY SeccionID, FechaSubida
        """
        
        self.cursor.execute(query, (incidente_id,))
        return self.cursor.fetchall()
    
    def _configurar_estilos(self, doc):
        """Configura los estilos del documento"""
        # Estilo para títulos
        titulo_style = doc.styles.add_style('TituloANCI', WD_STYLE_TYPE.PARAGRAPH)
        titulo_style.font.name = 'Arial'
        titulo_style.font.size = Pt(16)
        titulo_style.font.bold = True
        titulo_style.font.color.rgb = RGBColor(0, 0, 139)
        
        # Estilo para subtítulos
        subtitulo_style = doc.styles.add_style('SubtituloANCI', WD_STYLE_TYPE.PARAGRAPH)
        subtitulo_style.font.name = 'Arial'
        subtitulo_style.font.size = Pt(14)
        subtitulo_style.font.bold = True
        
        # Estilo para campos
        campo_style = doc.styles.add_style('CampoANCI', WD_STYLE_TYPE.PARAGRAPH)
        campo_style.font.name = 'Arial'
        campo_style.font.size = Pt(11)
    
    def _agregar_encabezado(self, doc, datos):
        """Agrega el encabezado del documento"""
        # Título principal
        titulo = doc.add_paragraph('FORMULARIO DE REPORTE DE INCIDENTE DE CIBERSEGURIDAD')
        titulo.style = 'TituloANCI'
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo
        subtitulo = doc.add_paragraph('AGENCIA NACIONAL DE CIBERSEGURIDAD E INFRAESTRUCTURA (ANCI)')
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Información básica
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        info_table.cell(0, 0).text = 'FOLIO ANCI:'
        info_table.cell(0, 1).text = str(datos.get('FolioANCI') or 'PENDIENTE')
        
        info_table.cell(1, 0).text = 'FECHA REPORTE:'
        info_table.cell(1, 1).text = datetime.now().strftime('%d/%m/%Y %H:%M')
        
        info_table.cell(2, 0).text = 'TIPO EMPRESA:'
        info_table.cell(2, 1).text = str(datos.get('TipoEmpresa') or 'N/A')
        
        info_table.cell(3, 0).text = 'ESTADO REPORTE:'
        info_table.cell(3, 1).text = str(datos.get('EstadoReporte') or 'EN PROCESO')
        
        doc.add_paragraph()
        
        # Información de la empresa
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('EMPRESA: ').bold = True
        p.add_run(str(datos.get('EmpresaNombre') or 'N/A') + ' - RUT: ' + str(datos.get('EmpresaRUT') or 'N/A'))
        doc.add_paragraph()
    
    def _agregar_seccion_1_reportante(self, doc, datos):
        """Agrega Sección 1: Datos del Reportante (No editable)"""
        # Agregar nota de sección protegida
        nota = doc.add_paragraph()
        nota.add_run('NOTA: Esta sección está protegida y no debe ser modificada.').italic = True
        nota.add_run(' ').font.size = Pt(9)
        nota.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_heading('SECCIÓN 1: DATOS DEL REPORTANTE', level=1)
        
        tabla = doc.add_table(rows=6, cols=2)
        tabla.style = 'Table Grid'
        
        # Campos del reportante
        campos = [
            ('Nombre completo:', str(datos.get('AnciNombreReportante') or '')),
            ('Cargo:', str(datos.get('AnciCargoReportante') or '')),
            ('Correo institucional:', str(datos.get('AnciCorreoReportante') or '')),
            ('Teléfono:', str(datos.get('AnciTelefonoReportante') or '')),
            ('Formación/Certificación:', str(datos.get('AnciFormacionCertificacion') or '')),
            ('Tipo de entidad:', str(datos.get('TipoEmpresa') or ''))
        ]
        
        for i, (campo, valor) in enumerate(campos):
            tabla.cell(i, 0).text = campo
            tabla.cell(i, 1).text = str(valor) if valor else 'No especificado'
            tabla.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        # Agregar nota al final de la sección 1
        doc.add_paragraph()
        nota_fin = doc.add_paragraph()
        nota_fin.add_run('--- FIN DE SECCIÓN PROTEGIDA ---').italic = True
        nota_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nota_fin.runs[0].font.size = Pt(9)
        nota_fin.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        doc.add_paragraph()
    
    def _agregar_seccion_2_identificacion(self, doc, datos):
        """Agrega Sección 2: Identificación del Incidente"""
        doc.add_heading('SECCIÓN 2: IDENTIFICACIÓN DEL INCIDENTE', level=1)
        
        # Información básica
        info_basica = [
            ('Fecha y hora de detección:', self._formatear_fecha(datos.get('FechaDeteccion'))),
            ('Fecha y hora de notificación:', self._formatear_fecha(datos.get('FechaGeneracion'))),
            ('Tipo de incidente (Taxonomía ANCI):', str(datos.get('AnciTipoIncidenteTaxonomia') or '')),
            ('Criticidad:', str(datos.get('Criticidad') or '')),
        ]
        
        for campo, valor in info_basica:
            p = doc.add_paragraph()
            p.add_run(f'{campo} ').bold = True
            p.add_run(str(valor) if valor else 'No especificado')
        
        # Descripción inicial
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Descripción inicial del incidente:').bold = True
        doc.add_paragraph(str(datos.get('DescripcionInicial') or 'Sin descripción disponible'))
        
        # Sistemas afectados
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Sistemas o servicios afectados:').bold = True
        doc.add_paragraph(str(datos.get('AnciSistemasAfectadosDetalle') or datos.get('SistemasAfectados') or 'No especificado'))
        
        # Impacto
        p = doc.add_paragraph()
        p.add_run('Impacto en continuidad o datos:').bold = True
        doc.add_paragraph(str(datos.get('AnciImpactoContinuidadDatos') or 'En evaluación'))
        
        # Afectación a terceros
        tabla_afectacion = doc.add_table(rows=2, cols=2)
        tabla_afectacion.style = 'Table Grid'
        
        tabla_afectacion.cell(0, 0).text = 'Afectación a terceros:'
        tabla_afectacion.cell(0, 1).text = 'SÍ' if datos.get('AnciAfectacionTerceros') else 'NO'
        
        tabla_afectacion.cell(1, 0).text = 'Número de afectados:'
        tabla_afectacion.cell(1, 1).text = str(datos.get('AnciNumeroAfectados') or 0)
        
        # Taxonomías ANCI
        if datos.get('taxonomias'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('Clasificación según Taxonomía ANCI:').bold = True
            
            tax_table = doc.add_table(rows=len(datos['taxonomias']) + 1, cols=4)
            tax_table.style = 'Table Grid'
            
            # Headers
            headers = ['Área', 'Efecto', 'Categoría', 'Descripción']
            for i, header in enumerate(headers):
                tax_table.cell(0, i).text = header
                tax_table.cell(0, i).paragraphs[0].runs[0].bold = True
            
            # Datos
            for i, tax in enumerate(datos['taxonomias'], 1):
                tax_table.cell(i, 0).text = str(tax[0]) if tax[0] else ''
                tax_table.cell(i, 1).text = str(tax[1]) if tax[1] else ''
                tax_table.cell(i, 2).text = str(tax[2]) if tax[2] else ''
                tax_table.cell(i, 3).text = str(tax[3]) if tax[3] else ''
        
        doc.add_paragraph()
    
    def _agregar_seccion_3_medidas(self, doc, datos):
        """Agrega Sección 3: Medidas Iniciales y Contención"""
        doc.add_heading('SECCIÓN 3: MEDIDAS INICIALES Y CONTENCIÓN', level=1)
        
        # Acciones de contención
        p = doc.add_paragraph()
        p.add_run('Acciones de contención aplicadas:').bold = True
        doc.add_paragraph(str(datos.get('AnciAccionesContencion') or 'Sin acciones registradas'))
        
        # Sistemas desconectados
        p = doc.add_paragraph()
        p.add_run('Sistemas desconectados o apagados:').bold = True
        doc.add_paragraph(str(datos.get('AnciSistemasDesconectados') or 'Ninguno'))
        
        # Notificaciones internas
        p = doc.add_paragraph()
        p.add_run('Notificaciones internas realizadas:').bold = True
        doc.add_paragraph(str(datos.get('AnciNotificacionesInternas') or 'Sin notificaciones'))
        
        # Estado actual
        p = doc.add_paragraph()
        p.add_run('Estado actual del incidente:').bold = True
        doc.add_paragraph(str(datos.get('AnciEstadoActualIncidente') or 'En investigación'))
        
        doc.add_paragraph()
    
    def _agregar_seccion_4_oiv(self, doc, datos):
        """Agrega Sección 4: Exclusiva para OIV"""
        doc.add_heading('SECCIÓN 4: INFORMACIÓN ADICIONAL PARA OIV', level=1)
        
        # SGSI
        tabla_oiv = doc.add_table(rows=9, cols=2)
        tabla_oiv.style = 'Table Grid'
        
        campos_oiv = [
            ('¿Cuenta con SGSI certificado (ISO 27001)?', 
             'SÍ' if datos.get('AnciOIVCuentaConSGSI') else 'NO'),
            ('Detalle Plan de Continuidad:', 
             str(datos.get('AnciOIVDetallePlanContinuidad') or '')),
            ('¿Se activó el plan de recuperación?', 
             str(datos.get('AnciOIVActivoPlanRecuperacion') or '')),
            ('Auditorías o ejercicios realizados:', 
             str(datos.get('AnciOIVAuditoriasRealizadas') or '')),
            ('Fecha última auditoría:', 
             self._formatear_fecha(datos.get('AnciOIVFechaUltimaAuditoria'))),
            ('Delegado técnico ante ANCI:', 
             str(datos.get('AnciOIVDelegadoTecnico') or '')),
            ('Medidas para evitar propagación:', 
             str(datos.get('AnciOIVMedidasPropagacion') or '')),
            ('Notificación a afectados:', 
             str(datos.get('AnciOIVNotificacionAfectados') or '')),
            ('Registro de capacitaciones:', 
             str(datos.get('AnciOIVRegistroCapacitaciones') or ''))
        ]
        
        for i, (campo, valor) in enumerate(campos_oiv):
            tabla_oiv.cell(i, 0).text = campo
            tabla_oiv.cell(i, 1).text = str(valor) if valor else 'No especificado'
            tabla_oiv.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
    
    def _agregar_anexos(self, doc, datos):
        """Agrega la sección de anexos y evidencias"""
        doc.add_page_break()
        doc.add_heading('ANEXOS: EVIDENCIAS Y DOCUMENTACIÓN', level=1)
        
        if datos.get('evidencias'):
            doc.add_paragraph('Listado de evidencias adjuntas:')
            
            evidencias_table = doc.add_table(rows=len(datos['evidencias']) + 1, cols=5)
            evidencias_table.style = 'Table Grid'
            
            # Headers
            headers = ['#', 'Sección', 'Archivo', 'Tamaño', 'Fecha']
            for i, header in enumerate(headers):
                evidencias_table.cell(0, i).text = header
                evidencias_table.cell(0, i).paragraphs[0].runs[0].bold = True
            
            # Datos
            for i, evidencia in enumerate(datos['evidencias'], 1):
                evidencias_table.cell(i, 0).text = str(i)
                evidencias_table.cell(i, 1).text = str(evidencia[0]) if evidencia[0] else ''
                evidencias_table.cell(i, 2).text = str(evidencia[1]) if evidencia[1] else ''
                evidencias_table.cell(i, 3).text = f"{str(evidencia[2]) if evidencia[2] else '0'} KB"
                evidencias_table.cell(i, 4).text = self._formatear_fecha(evidencia[3])
        else:
            doc.add_paragraph('No se han adjuntado evidencias a este reporte.')
        
        # Documentos requeridos
        doc.add_paragraph()
        doc.add_heading('Documentos adjuntos requeridos:', level=2)
        doc.add_paragraph('☐ Nombramiento del reportante con firma electrónica avanzada')
        doc.add_paragraph('☐ Registro del incidente (PDF o digital)')
        doc.add_paragraph('☐ Evidencia técnica (logs, capturas, informes)')
        if datos.get('TipoEmpresa') == 'OIV':
            doc.add_paragraph('☐ Certificados SGSI / Continuidad')
        doc.add_paragraph('☐ Bitácora de medidas tomadas')
    
    def _agregar_contacto_seguimiento(self, doc, datos):
        """Agrega la sección de contacto para seguimiento"""
        doc.add_paragraph()
        doc.add_heading('CONTACTO PARA SEGUIMIENTO DEL CASO', level=1)
        
        tabla_contacto = doc.add_table(rows=4, cols=2)
        tabla_contacto.style = 'Table Grid'
        
        campos_contacto = [
            ('Nombre y cargo:', str(datos.get('AnciContactoSeguimientoNombre') or '')),
            ('Horario disponible:', str(datos.get('AnciContactoSeguimientoHorario') or '')),
            ('Correo de emergencia:', str(datos.get('AnciContactoSeguimientoCorreo') or '')),
            ('Teléfono directo:', str(datos.get('AnciContactoSeguimientoTelefono') or ''))
        ]
        
        for i, (campo, valor) in enumerate(campos_contacto):
            tabla_contacto.cell(i, 0).text = campo
            tabla_contacto.cell(i, 1).text = str(valor) if valor else 'No especificado'
            tabla_contacto.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        # Nota final
        doc.add_paragraph()
        doc.add_paragraph()
        nota = doc.add_paragraph()
        nota.add_run('NOTA IMPORTANTE: ').bold = True
        nota.add_run(
            'Este formulario debe ser completado y enviado a través de la plataforma '
            'oficial de la ANCI. Los plazos de notificación corren desde el momento '
            'de la detección del incidente.'
        )
        
        # Firma
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('_' * 50)
        doc.add_paragraph('Firma del Reportante')
        doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
    
    def _formatear_fecha(self, fecha):
        """Formatea una fecha para mostrar en el informe"""
        if not fecha:
            return 'No especificada'
        
        try:
            if isinstance(fecha, str):
                fecha = datetime.fromisoformat(fecha.replace('Z', '+00:00'))
            return fecha.strftime('%d/%m/%Y %H:%M')
        except:
            return str(fecha)

# Función auxiliar para ser llamada desde otros módulos
def generar_informe_anci_completo(reporte_id):
    """
    Función helper para generar un informe ANCI completo
    
    Args:
        reporte_id: ID del reporte ANCI
        
    Returns:
        Ruta del archivo generado
    """
    generador = GeneradorInformeANCICompleto()
    return generador.generar_informe_docx(reporte_id)